from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches
from docx.shared import Pt
import os

def Data(H,CObj,COutcomes,partA): #to create template
 
        d= Document() #document docx object
    #--------------overall margin--------------------------------------
        sections=d.sections #sections is used here to set the margin size
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
#HEADER-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------Header------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
        table=d.add_table(5,4)
        table.style="Table Grid"
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for i in range(5):
            for j in range(4):
                table.cell(i,j).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c1=table.cell(0,1)
        c3=table.cell(0,3)
        c10=table.cell(1,0)
        c12=table.cell(1,2)
        c13=table.cell(1,3)
        table.cell(0,0).merge(c1)
        table.cell(0,2).merge(c3)
        table.cell(1,1).merge(c12)
        table.cell(1,1).merge(c13)
        paragraph = c10.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(os.getcwd() + '\img\CIT Logo Simple - PNG.png')
        p11=table.cell(1,1).add_paragraph(text="CHENNAI INSTITUTE OF TECHNOLOGY\n Sarathy Nagar, Pudupedu, Chennai - 600 069."+"\n"+str(H[0])+"-"+str(H[1])+" "+str(H[2])+" "+str(H[3]))
        p11.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p11=table.cell(1,1).add_paragraph(text=" ")
        table.cell(0,0).text="                                                                                                                      Reg. No.:"
        table.cell(2,0).text="Date /Time"
        table.cell(2,1).text=str(H[4])
        table.cell(2,2).text="Max. Marks"
        table.cell(2,3).text=str(H[9])
        table.cell(3,0).text="Subject with Code"
        table.cell(3,1).text=str(H[5])+ "-" + str(H[6])
        table.cell(3,2).text="Time"
        table.cell(3,3).text=str(H[10])
        table.cell(4,0).text="Branch"   #str(H[7])
        table.cell(4,1).text=str(H[7])+" - SET "+str(H[8])
        table.cell(4,2).text="Year/Semester"
        table.cell(4,3).text=str(H[11])+"/"+str(H[12])
        table.cell(2,1).width = Inches(5.0)
#COURSE OBJECTIVES--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------Course Objectives-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
        d.add_paragraph("Course Objectives\nThe Student should be able ")

         #course objective size is a variable
        tableCObj=d.add_table(len(CObj)+1,2)
        # for _ in range(cobj+1):
        #     CourseObjectives.append([])
        S_no=tableCObj.cell(0,0).add_paragraph(text="Sl.no")
        S_no.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        Course_Objectives_Title=tableCObj.cell(0,1).add_paragraph(text="Course Objectives\n")
        Course_Objectives_Title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for i in range(1,len(CObj)+1):
                tableCObj.cell(i,0).text="               "+str(i)
                tableCObj.cell(i,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for i in range(1,len(CObj)+1):
                tableCObj.cell(i,1).text=CObj[i-1]

          
        tableCObj.cell(0,1).width=Inches(22.0)
        #table.cell(0,0).width=Inches(5)
        tableCObj.style="Table Grid"
#COURSE OUTCOMES---------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------Course Outcomes----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
        d.add_paragraph("Course Outcomes:\nOn Completion of the course the students will be able  ")

        tableCOut=d.add_table(len(COutcomes[0])+1,2)

        C_no=tableCOut.cell(0,0).add_paragraph(text="CO No.")
        C_no.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        Course_Outcomes_Title=tableCOut.cell(0,1).add_paragraph(text="Course Outcomes\n")
        Course_Outcomes_Title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for i in range(1,len(COutcomes[0])+1):
              tableCOut.cell(i,0).text="            "+COutcomes[0][i-1]
              tableCOut.cell(i,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for i in range(1,len(COutcomes[0])+1):
              tableCOut.cell(i,1).text=COutcomes[1][i-1]
        tableCOut.cell(0,1).width=Inches(20.0)

        #table.cell(0,0).width=Inches(5)
        tableCOut.style="Table Grid"
#PART A ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------PartA--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
        d.add_paragraph("\n")

        table1a = d.add_table(1,1)
        table1a.style="Table Grid"
        for i in range(1):
            for j in range(1):
                table1a.cell(i,j).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER   

        #INSERTING AND STYLING CONTENTS TO TABLE 1
        p1=table1a.cell(0,0).add_paragraph(text="BLOOMS TAXONOMY(BT)\nK1-Remembering , K2-Understanding, K3-Applying, K4-Analyzing, K5-Evaluating ,K6-Creating\nN/D-November/December A/M-April/June\n")
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # p1bold=p1.add_run()
        # p1bold.bold=True
        #CREATING AND STYLING TABLE 2
        table2a = d.add_table(1,4)
        table2a.style = "Table Grid"
        table2a.cell(0,0).width = Inches(20.0)

        #INSERTING AND STYLING CONTENTS TO TABLE 2
        if H[9]=="50 Marks":
            p2=table2a.cell(0,0).add_paragraph(text="Part A-(5x2=10 marks)\n (Answer all the questions)\n")
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif H[9]=="100 Marks":
            p2=table2a.cell(0,0).add_paragraph(text="Part A-(10x2=20 marks)\n (Answer all the questions)\n")
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        p3=table2a.cell(0,1).add_paragraph(text="CO")
        p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p4=table2a.cell(0,2).add_paragraph(text="BT level")
        p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p5=table2a.cell(0,3).add_paragraph(text="Univ Qp \n reference")
        p5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        #CREATING AND STYLING TABLE 3
        if H[9]=="50 Marks":
           table3a = d.add_table(5,5)
           table3a.style = "Table Grid"
        elif H[9]=="100 Marks":
            table3a = d.add_table(10,5)
            table3a.style = "Table Grid"

        #INSERTING AND STYLING CONTENTS TO TABLE 3
        table3a.cell(0,1).width = Inches(15.0)
        qns=0
        if H[9]=="50 Marks":
             qns=5
        elif H[9]=="100 Marks":
             qns=10
        for i in range(qns):
            table3a.cell(i,0).add_paragraph(text="{}".format(i+1)).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for i in range(qns):
            table3a.cell(i,1).text=partA[0][i]
            table3a.cell(i,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            #questions
        for i in range(qns):
            table3a.cell(i,2).add_paragraph(text=partA[1][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #CO
        for i in range(qns):
            table3a.cell(i,3).add_paragraph(text=partA[2][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #BT Level
        for i in range(qns):
            table3a.cell(i,4).add_paragraph(text=partA[3][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #Univ QP reference


        d.add_paragraph("\n\n\n")
        d.add_paragraph("Prepared by                                                                 Verified by                                                                 Approved by").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        cwd = os.getcwd() + "/resrc/"

        d.save(cwd + H[5]+" "+H[8]+".docx")
    
        os.startfile(cwd + H[5]+" "+H[8]+".docx")