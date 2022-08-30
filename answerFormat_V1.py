from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches
from docx.shared import Pt
import os

class Header:
    #exam type - Internal Assessment, Model Exam, etc
    exam=["Internal Assessment","Model Exam"]
    MONTH=[0,"January","February","March","April","May","June","July","August","September","October","November","December"]
    branch=["CSE","EEE","ECE","IT","MECH","MCT","CSBS","AIDS","BME","CVL","S&H"]
    Set=["A","B"]
    marks=["50 Marks","100 Marks"]
    time=["1.30 hours","3 hours"]
    AcaYear=[0,"I","II","III","IV"]
    Sem=[0,"I","II","III","IV","V","VI","VII","VIII"]
    H=[]
    def __init__(self,exam,examno,month,year,date,subjcode,subjname,branch,Set,marks,AcaYear,Sem):
        self.H.append(self.exam[exam])#int 0-1
        self.H.append(examno) 
        self.H.append(self.MONTH[month])#int val
        self.H.append(year)#int val
        self.H.append(str(date)+"/"+str(month)+"/"+str(year))
        self.H.append(subjcode)#str
        self.H.append(subjname)#str
        self.H.append(self.branch[branch])#int 0 - 10
        self.H.append(self.Set[Set]) #int 0-1
        self.H.append(self.marks[marks]) #int 0-1
        self.H.append(self.time[marks])
        self.H.append(self.AcaYear[AcaYear]) #int 1-4
        self.H.append(self.Sem[Sem]) # int 1-8
   
def answerKey(exam,examno,month,year,date,subjcode,subjname,branch,Set,marks,AcaYear,Sem,a_list):
    a = Header(exam,examno,month,year,date,subjcode,subjname,branch,Set,marks,AcaYear,Sem)
    d= Document()
    sections=d.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    table=d.add_table(4,4)
    table.style="Table Grid"
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for i in range(4):
        for j in range(4):
              table.cell(i,j).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
    c1=table.cell(0,1)
    c2=table.cell(0,2)
    c3=table.cell(0,3)
    c10=table.cell(0,0)
    c12=table.cell(1,2)
    c13=table.cell(1,3)
    table.cell(0,1).merge(c2)
    table.cell(0,2).merge(c3)
    paragraph = c10.paragraphs[0]
    table.cell(0,1).width = Inches(10.0)
    
    paragraph = c10.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(os.getcwd() + '\img\CIT Logo Simple - PNG.png')
    
    p11=table.cell(0,1).add_paragraph()
    run=p11.add_run('CHENNAI INSTITUTE OF TECHNOLOGY\n')
    f=run.font
    f.name='Times New Roman'
    f.size=Pt(22)
    run=p11.add_run('Sarathy Nagar, Pudupedu, Chennai– 600 069.\nAssessment: I April  2022')
    f=run.font
    f.name='Times New Roman'
    f.size=Pt(17)
    p11.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p11=table.cell(1,1).add_paragraph(text=" ")
    
    
    
    table.cell(1,0).text="Date/Time"
    table.cell(1,1).text=a.H[4]
    table.cell(1,2).text="Max. Marks"
    table.cell(1,3).text=a.H[9]
    table.cell(2,0).text="Subject with Code"
    table.cell(2,1).text=a.H[6] + '-' + a.H[5]
    table.cell(2,2).text="Time"
    table.cell(2,3).text=a.H[10]
    table.cell(3,0).text="Branch  "
    table.cell(3,1).text=a.H[7]
    table.cell(3,2).text="Year/Semester"
    table.cell(3,3).text=a.H[11] + '/' + a.H[12]
    table.cell(1,1).width = Inches(5.0)
    
    a= d.add_paragraph()
    r=a.add_run('\n\nAnswer Key')
    a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r.underline=True
    f=r.font
    f.size=Pt(16)
    f.name='Times New Roman'
    
    a= d.add_paragraph()
    r=a.add_run('\n\nPart-A')
    r.underline=True
    f=r.font
    f.size=Pt(13)
    f.name='Times New Roman'
    
    #CREATING AND STYLING TABLE 3
    table3 = d.add_table(5,2)

    # partA=["Write the charcteristics of an algoritm.","How to measure algorithm running time?","Write the non-recursive algorithm to find the largest element in the list of numbers.","What is mean by Algorithm Visualization?","Define recurrance relation."]
    #INSERTING AND STYLING CONTENTS TO TABLE 3
    for i in range(5):
        table3.cell(i,0).width = Inches(1.0)
        table3.cell(i,1).width = Inches(16.0)
        table3.cell(i,0).text = "    {}".format(i+1)
        table3.cell(i,1).text = a_list[i]

    d.save(os.getcwd() + '/resrc/anskey.docx')

    os.startfile(os.getcwd() + '/resrc/anskey.docx')
    
# answerKey(0,1,4,2022,7,"CS000","DAA",0,0,0,2,4)